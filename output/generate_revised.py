#!/usr/bin/env python3
"""Generate professionally revised DOCX for multiplex hypergraph SIR paper."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

IMG_DIR = '/tmp/claude-0/-home-user-Claude-skills/c3202eb0-ca51-5b80-8b90-73153c19a107/scratchpad/docx_images'

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Heading styles
for level, (sz, name_cn) in enumerate([(Pt(16), '黑体'), (Pt(14), '黑体'), (Pt(13), '黑体')], 1):
    hs = doc.styles[f'Heading {level}']
    hs.font.size = sz
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.name = name_cn
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)

def add_para(text, style='Normal', bold=False, italic=False, size=None, alignment=None, space_after=Pt(6), first_line_indent=None):
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = size
    return p

def add_body(text):
    return add_para(text, first_line_indent=Cm(0.74))

def add_ref_item(text):
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.left_indent = Cm(0.74)
    pf.first_line_indent = Cm(-0.74)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_figure(img_name, caption, width=Inches(5.8)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    img_path = os.path.join(IMG_DIR, img_name)
    r.add_picture(img_path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cr = cap.add_run(caption)
    cr.font.size = Pt(10.5)
    cr.italic = True

# ═══════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════
add_para('多重超图上的 SIR 传播动力学：', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(18), space_after=Pt(2))
add_para('基于消息传递群体闭合的爆发阈值分析', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(18), space_after=Pt(18))

# ═══════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════
add_para('摘　要', bold=True, size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))

add_body(
    '传染病在人群中的传播机制长期以来是数学流行病学的核心课题 [1,2]。'
    '经典的 Kermack–McKendrick 均匀混合假设忽略了个体间接触结构的异质性 [3]，'
    '而基于网络的传播模型通过将宿主建模为图的顶点、将可能的传播途径建模为边，'
    '显著提升了对现实传播过程的刻画能力 [4,5]。'
    '然而，现有的网络传染病模型大多局限于成对（pairwise）交互 [6]，'
    '未能刻画多人群体中因协同暴露而产生的高阶传播效应。'
    '本文在多重超图（multiplex hypergraph）框架下研究 SIR 传播动力学，'
    '发展了一套精确的消息传递群体闭合方法。'
    '与有向超图上的逐成员乘积闭合不同，无向群体中群内级联（intra-group cascade）使得被感染成员延长群体的激活时长，'
    '传统的逐成员独立性假设因此失效。'
    '本文通过将群体内全部成员的计数状态 (s,i,r) 纳入边基变量（edge-based variable），'
    '在局域树状（locally tree-like）假设下导出了一组与系统规模无关的闭合常微分方程组，'
    '并由此给出显式的多型分支过程次代矩阵（next-generation matrix）及爆发阈值判据 ρ(N)=1。'
    '通过亚临界簇规模外推、精确 Gillespie 仿真及多重独立校验，'
    '在八组结构构型上确认了闭合解与仿真的一致性（最大相对偏差 6.5×10⁻⁴）。'
    '进一步的结构分析揭示了三个关键发现：'
    '(i) 层间参与相关（inter-layer degree correlation）可移动阈值达 12.5%；'
    '(ii) 多条弱渠道的协同效应可使合并后的阈值比任一单层阈值低 68%；'
    '(iii) 层间群体重叠（inter-layer overlap）可移动阈值达 115%，'
    '定量标定了树状闭合的适用边界。'
)

add_para('关键词：网络传染病学；多重超图；高阶交互；SIR 模型；消息传递；群体闭合；爆发阈值；生成函数',
         bold=True, size=Pt(10.5), space_after=Pt(18))

# ═══════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════
doc.add_heading('1　引言', level=1)

add_body(
    '自 Bernoulli 于 1760 年代建立天花接种的微分方程模型以来 [1]，'
    '数学模型在传染病传播研究中发挥着不可替代的作用。'
    '1927 年 Kermack 与 McKendrick 提出的 SIR 仓室模型奠定了现代数学流行病学的基础 [2]，'
    '其核心假设——均匀混合（mass-action mixing）——意味着任意两个个体间的接触概率相同。'
    '然而，大量实证研究表明，真实社会中的接触模式具有显著的异质性和结构性 [7,8]：'
    '个体的接触数服从重尾分布，高度数节点充当超级传播者（superspreader），'
    '群体内的聚集效应和社区结构深刻影响着传播动力学。'
)

add_body(
    '为突破均匀混合的局限，Watts 与 Strogatz [9] 以及 Barabási 与 Albert [10] '
    '分别提出了小世界网络和无标度网络模型，开启了网络流行病学的研究范式。'
    'Newman [4] 将渗流理论（percolation theory）引入网络传染病模型，'
    '建立了传播概率生成函数（probability generating function, PGF）框架；'
    'Allard 等 [5] 进一步将该框架推广至有向图，'
    '系统阐明了方向性（directionality）、异质性（heterogeneity）与相关性（correlation）'
    '在传播风险（risk）与传播能力（spread）中的不同角色。'
    '特别地，Allard 等 [5] 指出有向图中存在两种截然不同的“友谊悖论”'
    '（forward and backward friendship paradox）：'
    '沿出度方向追踪倾向于找到高风险个体，沿入度方向则倾向于找到高传播能力个体，'
    '这对接触追踪策略（contact tracing strategy）具有重要的政策启示。'
)

add_body(
    '上述工作虽然大幅推进了网络传染病学的理论深度，'
    '但几乎全部局限于成对交互的框架——每条边连接恰好两个节点。'
    '现实中许多传播场景本质上涉及多人群体的高阶交互（higher-order interaction）：'
    '家庭、教室、办公室、聚会等场合中的传播并非独立的两两接触之叠加，'
    '而是由群体内的协同暴露驱动 [11,12]。'
    '近年来，超图（hypergraph）和单纯复形（simplicial complex）'
    '作为刻画高阶交互的数学工具受到广泛关注 [13,14]。'
    'Ding 与 Zhu [15] 在含有向边的高阶单纯复形网络上构建了信息-疫情双层耦合模型，'
    '利用微观 Markov 链方法（MMCA）推导了传播阈值并分析了 Turing 不稳定性，'
    '展示了高阶结构对传播动力学的深刻影响。'
)

add_body(
    '然而，现有的高阶网络传染病模型面临两个关键挑战。'
    '其一，多数工作采用均场近似或成对近似（pair approximation, PA），'
    '无法精确处理群体内部的级联效应——'
    '当一个群体中有成员被感染后，群体持续激活，'
    '后续被感染的成员进一步延长激活时长，形成自增强的群内级联。'
    '其二，多层网络（multiplex network）中不同传播渠道之间的协同与竞争 [16,17] '
    '尚未在高阶框架下得到系统处理。'
    '本文正是针对这两个挑战，在多重超图上发展精确的消息传递闭合方法。'
)

add_body(
    '本文的主要贡献如下。'
    '(1) 提出了多重超图上 SIR 传播的群体闭合方程组，'
    '其维数仅与各层群体基数有关而与系统规模 N 无关，'
    '可精确复现整个传播过程的含时演化。'
    '(2) 导出了显式的多型次代矩阵，给出爆发阈值的谱半径判据 ρ(N)=1，'
    '并通过亚临界簇规模外推提供了完全独立的非线性验证。'
    '(3) 系统分析了层间参与相关、渠道配置、群体粒度和层间重叠等结构自由度对阈值的影响，'
    '其中层间重叠的可证伪检验定量标定了树状闭合的适用边界。'
    '全文结构安排如下：第 2 节建立模型并导出闭合方程组；'
    '第 3 节给出次代矩阵与爆发阈值，并通过亚临界外推进行独立验证；'
    '第 4 节分析阈值的结构依赖性与理论的边界；第 5 节总结全文。'
)

# ═══════════════════════════════════════════════════════
# 2. MODEL
# ═══════════════════════════════════════════════════════
doc.add_heading('2　模型与群体闭合方程组', level=1)

doc.add_heading('2.1　多重超图的结构表示', level=2)

add_body(
    '本文考虑的传播基底为多重超图（multiplex hypergraph）K=(V,E,Δ₂)，'
    '其中 V 为节点集（|V|=N），E 为超边集，Δ₂ 为层标记。'
    '每层 a=1,…,M 对应一种传播渠道，层 a 是 V 上的超图 Hₐ=(V,Eₐ)，'
    '主线分析取同质基数 |e|=mₐ。'
    '节点 v 在层 a 中的层度（layer degree）定义为 k⁽ᵃ⁾(v)=|{e∈Eₐ:v∈e}|，'
    '写成层度向量 k(v)=(k⁽¹⁾,…,k⁽ᴹ⁾)。'
    '联合层度分布 P(k) 是本模型的结构输入，'
    '它编码了不同传播渠道之间的参与相关性（participation correlation）。'
    '如图 1(a) 所示，K₁=8 个节点的网络包含两层——层 1 由基数 m₁=3 的三角形群体组成，'
    '层 2 由基数 m₂=4 的四元群体组成——节点 u 的层度向量为 k(u)=(2,1)。'
)

add_body(
    '引入两个概率生成函数（probability generating function, PGF）[4,5,18]：')

add_para('\tΨ(x) = Σₖ P(k) ∏_c x_c^{k_c}\t(2.1)', space_after=Pt(4))
add_para('\tψₐ(x) = (1/⟨k⁽ᵃ⁾⟩) Σₖ kₐ P(k) ∏_c x_c^{k_c − δ_{ac}}\t(2.2)', space_after=Pt(4))

add_body(
    'Ψ 为节点侧生成函数，描述一个随机选取节点的层度分布；'
    'ψₐ 为余度生成函数（excess degree PGF），描述沿层 a 的一条超边到达某节点时'
    '该节点在各层的余度分布 [4,5]。'
    '指数中的 −δ_{ac} 正是余度相减：来路群体不再计入，'
    '这是网络渗流理论中的标准处理 [4]。'
    '层间的一切耦合效应最终只经由 Ψ 和 ψₐ 的交叉项进入动力学方程——'
    '这一性质使得层间参与相关对阈值的影响可以被显式追踪（§4.1）。'
)

doc.add_heading('2.2　SIR 传播动力学', level=2)

add_body(
    '节点状态 x_v ∈ {S, I, R}（易感、感染、康复）跨层共享——'
    '个体在任一渠道被感染即在所有渠道中都转为感染态，'
    '这与经典 SIR 模型中的非可逆性（non-recurrent dynamics）一致 [5]。'
    '记层 a 的群体 e 中已感染成员数为 nₑ(t)=|e∩I(t)|。'
    '引入群体激活阈值 θₐ：当 nₑ≥θₐ 时群体被激活，'
    '向 e 中每个易感成员各以速率 βₐ 独立传递感染，否则静默。'
    '感染者以速率 μ 康复。'
    '易感节点 u 的总风险率（hazard rate）为'
)

add_para('\tΛ_u(t) = Σₐ Σ_{e∈Eₐ: u∈e} βₐ 𝟙[nₑ(t)≥θₐ]\t(2.3)', space_after=Pt(4))

add_body(
    '以 μ 归一化时间后，无量纲传播速率 λₐ=βₐ/μ 成为控制参数，'
    '主线分析取 θₐ=1。'
    '与有向超图不同 [15,19]，无向群体中每个成员既是传染源（source）也是传染靶（target）：'
    '一个成员被感染后群体激活，可感染同群其他成员，'
    '后者被感染又延长群体的激活时长（active period），形成群内级联（intra-group cascade）。'
    '如图 1(b) 所示，级联使得群体的实际感染产出 C(m,λ,θ) 严格大于朴素估计 (m−1)T，'
    '其中 T=λ/(1+λ) 为成对传递概率（transmissibility）[4]。'
    '本模型的核心新内容正源于对群内级联的精确处理。'
)

doc.add_heading('2.3　空腔构造与边基变量', level=2)

add_body(
    '闭合方法的构造遵循消息传递（message passing）的标准范式 [4,20,21]。'
    '固定一个检验节点 u，令其永久易感（空腔节点，cavity node）。'
    '取 u 所属的一个层 a 群体 e，定义边基变量（edge-based compartmental variable）[22]：'
)

add_para('\tΦₐ(t) = Pr[e 到 t 时刻尚未向 u 传递感染]\t(2.4)', space_after=Pt(4))

add_body(
    '在局域树状假设下——位形模型（configuration model）系综在 N→∞ 时'
    '局域收敛于超树（hypertree），'
    '回路出现概率为 O(1/N) [4,5]——'
    'u 所属各群体的上游分支互不相交，节点侧因而因子化（factorize）为'
)

add_para('\tS(t) = (1−ε) Ψ(Φ₁(t),…,Φ_M(t))\t(2.5)', space_after=Pt(4))

add_body(
    '其中 ε 为均匀初始感染比例。'
    '以 e 除 u 外的 mₐ−1 个成员的计数 (s,i,r)（s+i+r=mₐ−1）为仓室，'
    '定义未传递子概率（sub-probability of non-transmission）：'
)

add_para('\tx⁽ᵃ⁾_{sir}(t) = Pr[e 尚未向 u 传递 ∧ e\\{u} 处于 (s,i,r)]\t(2.6)', space_after=Pt(4))

add_body(
    '则 Φₐ=Σ x⁽ᵃ⁾_{sir}，激活且尚未传递的部分为 Φ^A_a=Σ_{i≥θ} x⁽ᵃ⁾_{sir}。'
    '一个成员经由其他群体到 t 仍易感的概率为 (1−ε)ψₐ(Φ)，'
    '其外部风险率（external hazard rate）为'
)

add_para('\thₐ(t) = −(d/dt)ln[(1−ε)ψₐ(Φ)] = Σ_c β_c Φ^A_c (∂_c ψₐ(Φ))/(ψₐ(Φ))\t(2.7)',
         space_after=Pt(4))

add_body(
    '群体计数方程组（group compartmental ODE system）随即建立：'
    '易感成员的总感染率为外部风险与群内风险之和 hₐ+βₐ𝟙[i≥θₐ]，'
    '群体向 u 的传递以速率 βₐ𝟙[i≥θₐ] 将概率移出「未传递」类。据此：'
)

add_para('\tẋ⁽ᵃ⁾_{sir} = (hₐ+βₐ𝟙[i−1≥θₐ])(s+1)x⁽ᵃ⁾_{s+1,i−1,r}', space_after=Pt(0))
add_para('\t\t− (hₐ+βₐ𝟙[i≥θₐ])s x⁽ᵃ⁾_{sir}', space_after=Pt(0))
add_para('\t\t+ μ[(i+1)x⁽ᵃ⁾_{s,i+1,r−1} − i x⁽ᵃ⁾_{sir}] − βₐ𝟙[i≥θₐ]x⁽ᵃ⁾_{sir}\t(2.8)',
         space_after=Pt(4))

add_body(
    '对 (2.8) 求和，外部风险项、群内感染项和康复项各自守恒相消，'
    '仅余 Φ̇ₐ=−βₐΦ^A_a，与 (2.7) 自洽。'
    '闭合方程组的维数为 Σₐ C(mₐ+1,2)，与系统规模 N 无关——'
    '例如 M=2、m=(3,4) 时仅为 16 维。'
    '节点侧的 S(t) 由 (2.5) 代数给出，'
    '终态规模（final epidemic size）R(∞)=1−(1−ε)Ψ(Φ(∞)) 直接读出 [4,5]，'
    '二者均不增加方程维数。'
)

# ── Figure 1 ──
add_figure('rId7.png',
    '图 1　多重超图 SIR 群体闭合的机制示意。'
    '(a) 多重超图的结构：层 1 由三角形群体（m₁=3）组成，层 2 由四元群体（m₂=4）组成，'
    '节点 u 的层度向量 k(u)=(2,1)。'
    '(b) 群内级联机制：朴素估计 (m−1)T 假设各成员独立传递，'
    '而级联中被感染成员延长群体激活时长，'
    '实际感染产出 C(m,λ,θ)>(m−1)T。'
    '(c) 群体层面的分支过程（branching process）：'
    '来路群体（层 a）因余度相减 −δ_{ab} 出列，'
    '次代矩阵元 N_{ab}=C(X_{ab}−δ_{ab})。',
    width=Inches(6.0))

doc.add_heading('2.4　群内级联与次代矩阵', level=2)

add_body(
    '群内级联量 C 可通过孤立群体的递推精确计算。'
    '记 u(i,s) 为从 i 个感染者、s 个易感者出发此后新增感染数的期望，'
    '以 α_i=λs·𝟙[i≥θ] 为群体的总感染率（以康复率为单位），则 [22]：'
)

add_para('\tu(i,s) = [α_i/(α_i+i)]·[1+u(i+1,s−1)] + [i/(α_i+i)]·u(i−1,s)\t(2.9)',
         space_after=Pt(4))

add_body(
    '边界为 u(0,s)=u(i,0)=0，级联规模 C(m,λ,θ)=u(1,m−1)。'
    '当 m=2 时 C=λ/(1+λ)=T 即为成对传递概率（transmissibility），'
    '级联效应消失而回到经典网络渗流 [4]；'
    'm≥3 时 C>(m−1)T——例如 λ=1 时 m=3,4,5 分别超出 11.1%、21.5% 和 31.0%——'
    '表明高阶群体中的级联放大效应不可忽略。'
    '当激活阈值 θ≥2 时，单个种子无法激活群体，故 C=0。'
)

add_body(
    '分支过程（branching process）[4,5] 必须定义在群体层面而非节点层面——'
    '若在节点层面朴素统计直接传染关系，群内级联的后继将被重复计入。'
    '按感染经由哪一层传来分型，经由层 a 感染的节点其层 b 群体数的期望'
    '为 ⟨k⁽ᵃ⁾k⁽ᵇ⁾⟩/⟨k⁽ᵃ⁾⟩−δ_{ab}，故多型次代矩阵（next-generation matrix）[23] 为'
)

add_para('\tN_{ab} = C(m_b,λ_b,θ_b)·[⟨k⁽ᵃ⁾k⁽ᵇ⁾⟩/⟨k⁽ᵃ⁾⟩ − δ_{ab}]\t(2.10)',
         space_after=Pt(4))

add_body(
    '爆发阈值由谱半径判据（spectral radius criterion）ρ(N)=1 确定 [23]。'
    '该矩阵统一了单层与多层、成对与高阶的全部特例：'
    '所有 mₐ=2 时 C=T，(2.10) 回到多层网络 SIR 的已知阈值 [16]；'
    'M=1、m=2 时传递概率阈值 T_c=⟨k⟩/(⟨k²⟩−⟨k⟩)，'
    '恰为经典的 Newman 渗流阈值 [4]。'
    '此外 N 逐元非负，故 ρ(N)≥max_a N_{aa}——'
    '多重超图的阈值不高于任一单层的阈值 [16]。'
)

# ═══════════════════════════════════════════════════════
# 3. THRESHOLD VALIDATION
# ═══════════════════════════════════════════════════════
doc.add_heading('3　爆发阈值的独立验证', level=1)

doc.add_heading('3.1　含时演化的精确仿真校验', level=2)

add_body(
    '闭合方程组不仅给出阈值，也精确复现整个传播过程的含时演化。'
    '图 2(a) 给出 M=2、m=(3,4)、N=6000、λ=1.6λ_c 下'
    '感染比例 I(t) 的闭合解与精确 Gillespie 仿真 [24] 的对比。'
    '联合层度分布 P(k) 在 (2,2) 和 (3,3) 上各占一半，'
    'λ_c=0.0993，ε=0.02。'
    '群体闭合与精确仿真在 S(t)、I(t) 上的最大偏差分别为 0.0019 和 0.0006，'
    '小于图中线宽；感染峰高（epidemic peak）相差约 0.3%，'
    '峰位（peak timing）均在 t≈4.4。'
    '残余偏差随系统规模增大而单调趋零，'
    '是可由增大 N 消除的有限尺寸效应（finite-size effect），'
    '而非闭合层级的系统偏差。'
)

doc.add_heading('3.2　亚临界簇规模外推', level=2)

add_body(
    '为提供不经线性化的独立阈值验证，'
    '本文发展了一种基于亚临界簇规模发散的外推方法。'
    '阈值以下单个种子只引发有限簇（finite cluster），'
    '但其平均总规模随 λ→λ_c⁻ 无界增长 [4]。'
    '以多型分支过程计数，平均总规模为'
)

add_para('\tχ(λ) = 1ᵀ(I−N)⁻¹ v₀\t(3.1)', space_after=Pt(4))

add_body(
    '其中 v₀ 为种子的型分布。'
    'χ 的极点来自 (I−N)⁻¹，恰在 ρ(N)=1 处按 (1−ρ)⁻¹ 发散。'
    '其倒数 ε/R(∞)=1/χ(λ) 是一条在阈值处线性穿零的光滑函数 [4]：'
)

add_para('\tε/R(∞) ≃ c·(λ_c−λ),\tλ→λ_c⁻\t(3.2)', space_after=Pt(4))

add_body(
    '零点即 λ_c。'
    '该判据与 ρ(N)=1 不共享任何中间量：'
    '一侧将非线性方程组积分至终态并作线性拟合，'
    '另一侧对次代矩阵求谱半径并二分，'
    '二者仅共享模型定义。'
    '两条路线的一致构成对彼此的非平凡交叉检验（cross-validation），'
    '正如图 2(b)–(d) 所报告的。'
)

add_body(
    '有限种子偏置（finite-seed bias）需予修正。'
    'ε>0 时即便恰在 λ_c 处 R(∞) 仍为有限，'
    '零点被系统性推高。'
    '该偏置可由 ε→0 的 Richardson 外推消去 [25]。'
    '不作外推时零点偏高 0.37%–0.44%，外推后降至 0.031%–0.065%。'
)

# ── Figure 2 ──
add_figure('rId8.png',
    '图 2　群体闭合的两重独立验证。'
    '(a) 含时演化：感染比例 I(t) 的闭合解（虚线）与精确 Gillespie 仿真（实线）逐点吻合'
    '（N=6000, 400 次实现, λ=1.6λ_c），灰带为 95% 置信区间。'
    '(b) 阈值的亚临界测定：逆平均簇规模 ε/R(∞)=1/χ 随 λ/λ_c 线性趋零，'
    '其 x 截距即 λ_c（三组构型，虚线为线性外推，空心圆为 ρ(N)=1 预言）。'
    '(c) 八组构型的外推 λ_c 与谱值 ρ(N)=1 落在对角线上，最大相对偏差 6.5×10⁻⁴。'
    '(d) 偏离的标准化残差，最大者 1.51σ，全部落在 2σ 以内。',
    width=Inches(5.5))

doc.add_heading('3.3　数值结果', level=2)

add_body(
    '在窗口 λ∈[0.90, 0.995]λ_c 上取十个采样点，'
    '每点按上述方式求 ε→0 的 ε/R(∞) 并作线性拟合。'
    '八组构型的阈值覆盖 λ_c∈[0.08, 0.40]，'
    '含单层高阶超图、成对退化网络和双层多重超图。'
    '外推截距与 ρ(N)=1 解析值的相对偏差最大为 6.5×10⁻⁴，'
    '以回归标准误 σ 为单位，偏离最大 1.51σ，全部落在 1.6σ 之内。'
    '残余偏差源于线性律的有限窗口效应（finite-window curvature），'
    '而非闭合与谱方法之间的实质分歧。'
)

doc.add_heading('3.4　数值校验体系', level=2)

add_body(
    '本文建立了四层递进的校验体系，以确保结果的可靠性。'
    '(i) 内部自洽：方程组的求和规则残差达机器精度（10⁻¹⁷），'
    '全易感恒等式 x⁽ᵃ⁾_{m-1,0,0}=[(1−ε)ψₐ(Φ)]^{m-1} 的残差按 dt⁴ 收敛，'
    '与四阶 Runge–Kutta 的阶一致。'
    '(ii) 解析锚点：m=2 退化到经典成对传递概率 T=λ/(1+λ)，'
    '5-正则单层的 λ_c=1/3 逐位精确。'
    '(iii) 独立复算：闭合的亚临界终态与 (3.1) 的解析值不共享代码，'
    '相对差随 ε 单调趋零——ε=10⁻⁵ 时降至 1.1×10⁻⁵。'
    '(iv) 蒙特卡洛复核：级联递推 (2.9) 由单群体的直接 Gillespie 抽样独立验证，'
    '2×10⁵ 次实现的估计与递推值一致，偏离均在统计标准误以内。'
)

# ═══════════════════════════════════════════════════════
# 4. STRUCTURAL DEPENDENCE
# ═══════════════════════════════════════════════════════
doc.add_heading('4　阈值的结构依赖性与理论的适用边界', level=1)

add_body(
    '次代矩阵 (2.10) 将复杂的网络结构压缩为两个要素：'
    '联合层度分布的二阶矩和各层的群体基数。'
    '本节沿四条结构方向系统探测阈值的敏感性，'
    '并通过一条公式看不见的自由度定量标定理论的适用边界。'
)

doc.add_heading('4.1　协同效应：弱渠道合并后的超临界', level=2)

add_body(
    '令两层的传播速率 λ₁ 与 λ₂ 独立变化，'
    '阈值条件 ρ(N)=1 在 (λ₁,λ₂) 平面上描出一条临界曲线。'
    '与之对照的是两条单层阈值线 N_{aa}=1。'
    '三条线围出一个楔形区域——协同区（synergy region）：'
)

add_para('\tS = {(λ₁,λ₂) : N₁₁<1, N₂₂<1, ρ(N)>1}\t(4.1)', space_after=Pt(4))

add_body(
    '楔内每一点对应同一个局面：任一渠道单独存在均不足以引发爆发，'
    '两者并存却足够——这是多层网络传播中协同效应（synergistic effect）[16,17] '
    '在高阶框架下的精确定量体现。'
    '如图 3(a) 所示，以 P={(2,2),(3,3)}、m=(3,3) 的构型为例，'
    '单层阈值 λ_c⁽ᵃ⁾=0.400567，两层对称合并后降至 0.128162，低 68.0%。'
    '协同区并非临界线附近的一条窄缝，'
    '而是占据了整个 [0,λ_c⁽ᵃ⁾]² 方块中临界曲线以上的全部区域。'
)

doc.add_heading('4.2　层间参与相关的效应', level=2)

add_body(
    '阈值只经由度矩进入 N，故可构造一族边缘分布固定、仅相关不同的构型，'
    '隔离出层间参与相关（inter-layer degree correlation）的纯粹效应。'
    '设两层度各以等概率为 2 或 4，联合分布取 '
    'P(2,2)=P(4,4)=(1+ρ₁₂)/4、P(2,4)=P(4,2)=(1−ρ₁₂)/4，'
    '其中 ρ₁₂ 恰为两层度的 Pearson 相关系数。'
    'ρ₁₂ 由 −1 扫到 +1 时边缘不变，只有交叉矩 ⟨k⁽¹⁾k⁽²⁾⟩=9+ρ₁₂ 改变，'
    '阈值 λ_c 随 ρ₁₂ 单调下降，共约 12.5%（图 3(b)）。'
    '其物理机制清晰：正相关意味着同一节点在两层同时充当枢纽（hub），'
    '在边缘结构不变时集中跨层传播能力 [5]，从而压低阈值；'
    '反相关则抬高阈值——这一移动在图 3(c)–(d) 的相图中表现为临界曲线的整体内移。'
    '位形模型多重超图上的直接 Gillespie 仿真（N=2×10⁴）在六个 ρ₁₂ 上复现该曲线，'
    '偏离最大 1.9σ。'
)

# ── Figure 3 ──
add_figure('rId9.png',
    '图 3　爆发阈值的相图。'
    '(a) ρ(N) 在 (λ₁,λ₂) 平面的热力图（magma 色标），'
    '白色加粗线为临界曲线 ρ(N)=1，白色虚线为单层阈值 N_{aa}=1，'
    '二者围出的楔形即协同区 S。'
    '(b) 固定边缘分布的 {2,4} 族：精确 λ_c(ρ₁₂)（蓝线）与直接仿真的自助箱线。'
    '(c) 同族在 ρ₁₂=−1,0,+1 下的临界曲线：边缘钉死，仅曲线随相关内移。'
    '(d) Δλ_{2,c} 相对 ρ₁₂=0 的移动量（发散色标）。',
    width=Inches(5.5))

doc.add_heading('4.3　渠道配置与群体粒度', level=2)

add_body(
    '最劣配置（worst-case allocation）。'
    '固定总传播预算 Σₐwₐ 在两层间转移，最劣配置总是内点，'
    '而非将预算全部押在最强渠道上。'
    '纯配置令某层 λ_b=0，把 N 的非对角元归零，'
    '跨层项恰是多重结构的全部收益所在——'
    '这在 Allard 等 [5] 关于有向图多型分支矩阵的分析中已有类似的结构洞察。'
    '图 4(a) 显示对称构型 m=(3,3)、k=(3,3) 的最劣配置比最优纯配置低 29.3%。'
)

add_body(
    '群体粒度（group granularity）。'
    '“少数大群体”与“多数小群体”孰更危险，取决于固定什么（图 4(b)）。'
    '固定每节点群体数 k=4，阈值从 m=2 的 0.500 降至 m=8 的 0.044，'
    '大群体因级联的超线性放大效应（superlinear cascade amplification）而更危险。'
    '改为固定接触预算 k(m−1)=12，结论反转：阈值从 m=2 的 0.100 升至 m=7 的 0.148。'
    '反转源于单层阈值条件 (k−1)·C(m,λ_c,θ)=1——'
    '增大 m 换来的 C 的超线性增益被余度 k−1 的坍缩所抵消。'
    '极端情形 k=1（余度为零）时该层永远无法独立传播，'
    '否证了“大群体更危险”的朴素预期。'
)

doc.add_heading('4.4　层间重叠：一条可证伪的推论', level=2)

add_body(
    '以上分析均在公式可见的自由度上进行。'
    '更能标定理论边界的，是一条公式看不见的自由度。'
    '次代矩阵 (2.10) 只读取联合层度分布与群体基数，'
    '因此两个 P(k) 与 m 相同的多重超图必被赋予同一个 λ_c，'
    '无论两层的群体在节点上如何相对摆放。'
    '定义层间重叠（inter-layer overlap）[26]：'
)

add_para('\to = #{共处层1群体∧共处层2群体的节点对} / #{共处某层1群体的节点对}\t(4.2)',
         space_after=Pt(4))

add_body(
    '一个被重复使用的节点对闭合了一条 4-循环（4-cycle），'
    '正是局域树状假设所丢弃的短环结构 [4,5]。'
    '由此得到一条可证伪的推论（falsifiable prediction）：'
    '固定 P(k)、m 与 θ，只改 o，则 λ_c 必须不动。'
    '若测得阈值随 o 移动，移动的幅度即为树状闭合失效的定量刻度。'
)

add_body(
    '实验构造如下。取每节点层度 (2,2)、m=(3,3)。'
    '层 1 为位形模型；层 2 将层 1 中随机选出的一部分群体逐字复制，'
    '其余按残余度随机连线，确保全族共享同一个 P(k) 与 m，o 为唯一变动量。'
    '理论给出与 o 无关的水平线 3C(3,λ_c)=1，即 λ_c=0.18614；'
    'o=1 时层 2 成为层 1 的副本，阈值退化为 1·C(3,2λ_c)=1，即 λ_c=0.40974。'
    '两条参照线相差 120%。'
)

add_body(
    '在六个重叠水平上以亚临界外推测定阈值（N=2×10⁴），'
    '所得 λ_c 由 o=0 的 0.18065 单调升至 o=1 的 0.38809，高出 115%，'
    '而理论预言为水平线——推论被干净地证伪（图 4(c)）。'
    'o=1 处偏离水平线达 30σ，逐点递增也都显著（相邻两点之差最小 4.5σ）。'
    '层间重叠不是可忽略的高阶修正，'
    '而是比层间参与相关（12.5%）更强的阈值调控量——'
    '相差近一个数量级。'
)

add_body(
    '失效方向与机制一致：重叠使传播投向已可及的节点，'
    '同一条边被两层重复计入，而 (2.10) 按互不相交的分支计数，'
    '系统性高估分支能力、低估阈值。'
    '图 4(d) 给出小 λ 下的展开：理论按 3C(3,λ)≈6λ 计数，'
    'o=1 的实际过程只有 C(3,2λ)≈4λ，比值 3:2 即高估量级。'
    '树状闭合的适用边界由此定量给出：'
    '位形模型系综自动满足 o=O(1/N)，'
    '而任何具有实质层间重叠的真实系统都将落在公式之外 [4,5]。'
)

# ── Figure 4 ──
add_figure('rId10.png',
    '图 4　阈值的结构依赖性与理论的适用边界。'
    '(a) 渠道配置：固定总预算在两层间转移，圆点为最劣配置——均为内点。'
    '(b) 群体粒度：固定 k=4 时阈值随 m 下降（大群体更危险），'
    '固定 k(m−1)=12 时反升（小群体更危险），两种归一给出相反结论。'
    '(c) 层间重叠 o 的可证伪检验：理论预言 λ_c 与 o 无关（蓝线），'
    '实测单调上升至 o=1 的解析值（红虚线），误差棒按 √(χ²/dof) 标定。'
    '(d) 机制图解：(2.10) 所计分支因子 3C(3,λ) 高于 o=1 实际的 C(3,2λ)，'
    '两者穿过 1 处即各自阈值。',
    width=Inches(5.5))

# ═══════════════════════════════════════════════════════
# 5. CONCLUSION
# ═══════════════════════════════════════════════════════
doc.add_heading('5　结论与展望', level=1)

add_body(
    '本文在多重超图框架下建立了 SIR 传播的精确消息传递群体闭合理论。'
    '通过将群体内全部成员的计数状态纳入边基变量，'
    '克服了高阶群体中群内级联导致逐成员独立性假设失效的核心困难，'
    '导出了一组维数仅与群体基数相关而与系统规模无关的闭合常微分方程组。'
    '该方程组不仅给出显式的多型次代矩阵和爆发阈值判据 ρ(N)=1，'
    '也精确复现了整个传播过程的含时演化——'
    '与精确 Gillespie 仿真的最大偏差小于 0.2%。'
)

add_body(
    '结构分析揭示了多重超图传播的若干非直觉特性。'
    '(i) 协同效应使合并阈值可比任一单层阈值低 68%，'
    '即两条独立无法引发爆发的弱渠道合并后足以引发大规模流行。'
    '(ii) 渠道配置的最劣点总是内点，将预算集中于单一渠道反而更安全。'
    '(iii) 群体粒度的效应取决于比较基准的选取——'
    '固定群体参与数与固定接触预算给出相反结论，'
    '否证了“大群体更危险”的朴素预期。'
    '(iv) 层间重叠的可证伪检验定量标定了树状闭合的适用边界，'
    '重叠可移动阈值达 115%，比层间参与相关（12.5%）强近一个数量级。'
)

add_body(
    '本工作为若干方向的后续研究奠定了基础。'
    '在理论层面，将闭合方法推广至 SIS 等可逆动力学、'
    '异质激活阈值和时变网络结构，'
    '是自然的延伸方向。'
    '在方法层面，结合 Ding 与 Zhu [15] 的最优控制和参数辨识框架，'
    '可将群体闭合方程组嵌入物理信息神经网络（PINN）[27]，'
    '实现对真实疫情数据的高阶网络拟合。'
    '在应用层面，层间重叠的量化对真实多层传播网络'
    '（如家庭-学校-工作场所的多层接触结构 [28]）'
    '的建模具有直接的指导意义：'
    '高重叠的真实系统需要超越树状闭合的方法，'
    '而低重叠系统则可信赖本文的理论预测。'
)

# ═══════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════
doc.add_heading('参考文献', level=1)

refs = [
    '[1] D. Bernoulli, "Essai d\'une nouvelle analyse de la mortalité causée par la petite vérole," Mém. Math. Phys. Acad. Roy. Sci., Paris, 1766.',
    '[2] W.O. Kermack, A.G. McKendrick, "A contribution to the mathematical theory of epidemics," Proc. R. Soc. Lond. A, 115(772):700–721, 1927.',
    '[3] H.W. Hethcote, "The mathematics of infectious diseases," SIAM Rev., 42(4):599–653, 2000.',
    '[4] M.E.J. Newman, "Spread of epidemic disease on networks," Phys. Rev. E, 66(1):016128, 2002.',
    '[5] A. Allard, C. Moore, S.V. Scarpino, B.M. Althouse, L. Hébert-Dufresne, "The role of directionality, heterogeneity, and correlations in epidemic risk and spread," SIAM Rev., 65(2):471–492, 2023.',
    '[6] R. Pastor-Satorras, C. Castellano, P. Van Mieghem, A. Vespignani, "Epidemic processes in complex networks," Rev. Mod. Phys., 87(3):925–979, 2015.',
    '[7] A.-L. Barabási, "Network Science," Cambridge Univ. Press, 2016.',
    '[8] M.E.J. Newman, "Networks: An Introduction," Oxford Univ. Press, 2010.',
    '[9] D.J. Watts, S.H. Strogatz, "Collective dynamics of \'small-world\' networks," Nature, 393(6684):440–442, 1998.',
    '[10] A.-L. Barabási, R. Albert, "Emergence of scaling in random networks," Science, 286(5439):509–512, 1999.',
    '[11] I. Iacopini, G. Petri, A. Barrat, V. Latora, "Simplicial models of social contagion," Nat. Commun., 10:2485, 2019.',
    '[12] G. Ferraz de Arruda, M. Tizzani, Y. Moreno, "Phase transitions and stability of dynamical processes on hypergraphs," Commun. Phys., 4:24, 2021.',
    '[13] F. Battiston, G. Cencetti, I. Iacopini, et al., "Networks beyond pairwise interactions: structure and dynamics," Phys. Rep., 874:1–92, 2020.',
    '[14] S. Majhi, M. Perc, D. Ghosh, "Dynamics on higher-order networks: a review," J. R. Soc. Interface, 19:20220043, 2022.',
    '[15] Y. Ding, L. Zhu, "Epidemic dynamics in higher-order networks with directed edges," Commun. Nonlinear Sci. Numer. Simul., 162:110433, 2026.',
    '[16] F.D. Sahneh, C. Scoglio, P. Van Mieghem, "Generalized epidemic mean-field model for spreading processes over multilayer complex networks," IEEE/ACM Trans. Netw., 21(5):1609–1620, 2013.',
    '[17] M. De Domenico, C. Granell, M.A. Porter, A. Arenas, "The physics of spreading processes in multilayer networks," Nat. Phys., 12(10):901–906, 2016.',
    '[18] H.S. Wilf, "Generatingfunctionology," 3rd ed., A K Peters, 2006.',
    '[19] G. St-Onge, H. Sun, A. Allard, L. Hébert-Dufresne, G. Bianconi, "Universal nonlinear infection kernel from heterogeneous exposure on higher-order networks," Phys. Rev. Lett., 127:158301, 2021.',
    '[20] B. Karrer, M.E.J. Newman, "Message passing approach for general epidemic models," Phys. Rev. E, 82:016101, 2010.',
    '[21] J.P. Gleeson, "Binary-state dynamics on complex networks: pair approximation and beyond," Phys. Rev. X, 3:021004, 2013.',
    '[22] J.C. Miller, A.C. Slim, E.M. Volz, "Edge-based compartmental modelling for infectious disease spread," J. R. Soc. Interface, 9:890–906, 2012.',
    '[23] O. Diekmann, J.A.P. Heesterbeek, J.A.J. Metz, "On the definition and the computation of the basic reproduction ratio R₀," J. Math. Biol., 28(4):365–382, 1990.',
    '[24] D.T. Gillespie, "Exact stochastic simulation of coupled chemical reactions," J. Phys. Chem., 81(25):2340–2361, 1977.',
    '[25] L.F. Richardson, "The deferred approach to the limit," Phil. Trans. R. Soc. A, 226:299–361, 1927.',
    '[26] V. Nicosia, G. Bianconi, V. Latora, M. Barthelemy, "Nonlinear growth and condensation in multiplex networks," Phys. Rev. E, 90:042807, 2014.',
    '[27] M. Raissi, P. Perdikaris, G.E. Karniadakis, "Physics-informed neural networks," J. Comput. Phys., 378:686–707, 2019.',
    '[28] N. Perra, "Non-pharmaceutical interventions during the COVID-19 pandemic: a review," Phys. Rep., 913:1–52, 2021.',
]

for ref in refs:
    add_ref_item(ref)

# ═══════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════
output_path = '/tmp/claude-0/-home-user-Claude-skills/c3202eb0-ca51-5b80-8b90-73153c19a107/scratchpad/docx_work/revised_output.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
